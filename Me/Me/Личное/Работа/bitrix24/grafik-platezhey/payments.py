"""График платежей: чтение сделки, генерация графика и работа со счетами Битрикса.

Счёт = стандартная сущность «Счета» (entityTypeId=31). Наш модуль управляет ТОЛЬКО
теми счетами, которые создал сам (по сохранённым invoice_id + пометке xmlId). Чужие
счета на сделке не трогает. Оплату/просрочку ведёт штатный функционал счетов.
"""
import os
import base64
import tempfile
import requests
from datetime import date, datetime
from calendar import monthrange

from config import BITRIX_WEBHOOK, PAYMENT_DOC_TEMPLATE_ID
import db

PAYMENTS_MARKER = "[[PAYMENTS_TABLE]]"

# ─── Константы счетов (проверено на портале dm-tmn) ───
INVOICE_ETYPE = 31          # сущность «Счета»
INVOICE_CATEGORY = 2        # воронка «Общая»
STAGE_NEW = "DT31_2:N"      # Новый
STAGE_PAID = "DT31_2:P"     # Оплачен


def _call(method, params=None):
    """Вызов вебхука Битрикса (POST JSON). Бросает при ошибке API."""
    url = f"{BITRIX_WEBHOOK}/{method}.json"
    r = requests.post(url, json=params or {}, timeout=30)
    r.raise_for_status()
    data = r.json()
    if isinstance(data, dict) and data.get("error"):
        raise RuntimeError(f"{data.get('error')}: {data.get('error_description')}")
    return data.get("result")


# ──────────────────────────── Сделка ────────────────────────────

def deal_context(deal_id):
    """Контекст сделки для счетов: сумма, валюта, компания, ответственный, название."""
    d = _call("crm.deal.get", {"id": deal_id}) or {}
    if not d:
        raise RuntimeError(f"Сделка {deal_id} не найдена")
    return {
        "id": int(d["ID"]),
        "title": d.get("TITLE") or f"Сделка #{deal_id}",
        "opportunity": round(float(d.get("OPPORTUNITY") or 0), 2),
        "currency": d.get("CURRENCY_ID") or "RUB",
        "company_id": int(d["COMPANY_ID"]) if d.get("COMPANY_ID") else None,
        "assigned_by_id": int(d["ASSIGNED_BY_ID"]) if d.get("ASSIGNED_BY_ID") else None,
    }


# ──────────────────────────── Даты и генерация графика ────────────────────────────

def _parse_date(s):
    """'YYYY-MM-DD' (или ISO с временем) → date."""
    return datetime.fromisoformat(str(s)[:10]).date()


def add_months(d, n):
    """Сдвиг даты на n месяцев с корректным концом месяца (31 янв + 1 мес = 28/29 фев)."""
    m = d.month - 1 + int(n)
    y = d.year + m // 12
    m = m % 12 + 1
    day = min(d.day, monthrange(y, m)[1])
    return date(y, m, day)


def build_equal(total, advance, count, first_date, freq_months):
    """График «равными частями»: аванс — первая строка (часть суммы), остаток равными
    долями по (count-1) платежам. Копеечный «хвост» — на последнюю строку, чтобы Σ = total."""
    total = round(float(total), 2)
    advance = round(float(advance or 0), 2)
    count = int(count)
    freq = int(freq_months)
    fd = _parse_date(first_date)

    if count <= 1:
        return [{"seq": 1, "due_date": fd.isoformat(), "amount": total}]

    rows = []
    if advance > 0:
        rows.append({"seq": 1, "due_date": fd.isoformat(), "amount": advance})
        remaining = round(total - advance, 2)
        rest_count = count - 1
        start_seq = 2
    else:
        remaining = total
        rest_count = count
        start_seq = 1

    share = round(remaining / rest_count, 2)
    for i in range(rest_count):
        seq = start_seq + i
        rows.append({
            "seq": seq,
            "due_date": add_months(fd, (seq - 1) * freq).isoformat(),
            "amount": share,
        })

    diff = round(total - sum(r["amount"] for r in rows), 2)
    if diff:
        rows[-1]["amount"] = round(rows[-1]["amount"] + diff, 2)
    return rows


def build_full(total, first_date):
    """График «полная оплата»: одна строка на всю сумму."""
    return [{"seq": 1, "due_date": _parse_date(first_date).isoformat(),
             "amount": round(float(total), 2)}]


def validate_rows(rows, total):
    """Проверка строк перед созданием счетов. Возвращает текст ошибки или None."""
    if not rows:
        return "График пуст"
    for r in rows:
        if float(r["amount"]) <= 0:
            return f"Платёж №{r.get('seq')}: сумма должна быть больше нуля"
    dates = [_parse_date(r["due_date"]) for r in rows]
    if any(dates[i] > dates[i + 1] for i in range(len(dates) - 1)):
        return "Даты платежей должны идти по возрастанию"
    s = round(sum(float(r["amount"]) for r in rows), 2)
    if abs(s - round(float(total), 2)) > 0.01:
        return f"Сумма платежей ({s}) не совпадает с суммой сделки ({total})"
    return None


# ──────────────────────────── Счета: CRUD ────────────────────────────

def list_invoices_by_ids(ids):
    """Актуальное состояние наших счетов по id: {id: {stageId, opportunity, closedate, ...}}."""
    ids = [i for i in ids if i]
    if not ids:
        return {}
    res = _call("crm.item.list", {
        "entityTypeId": INVOICE_ETYPE,
        "filter": {"@id": ids},
        "select": ["id", "stageId", "opportunity", "closedate", "accountNumber"],
    }) or {}
    return {int(it["id"]): it for it in res.get("items", [])}


def create_invoice(deal_id, seq, amount, due_date, ctx):
    """Создать счёт под строку графика. Возвращает id счёта."""
    fields = {
        "title": f"Платёж {seq} — {ctx['title']}",
        "opportunity": amount,
        "currencyId": ctx.get("currency") or "RUB",
        "closedate": due_date,
        "begindate": date.today().isoformat(),
        "stageId": STAGE_NEW,
        "categoryId": INVOICE_CATEGORY,
        "parentId2": deal_id,
        "xmlId": f"paysched:{deal_id}:{seq}",
    }
    if ctx.get("company_id"):
        fields["companyId"] = ctx["company_id"]
    if ctx.get("assigned_by_id"):
        fields["assignedById"] = ctx["assigned_by_id"]
    res = _call("crm.item.add", {"entityTypeId": INVOICE_ETYPE, "fields": fields}) or {}
    return int(res["item"]["id"])


def update_invoice(invoice_id, amount, due_date):
    """Обновить сумму и срок неоплаченного счёта."""
    _call("crm.item.update", {
        "entityTypeId": INVOICE_ETYPE,
        "id": invoice_id,
        "fields": {"opportunity": amount, "closedate": due_date},
    })


def delete_invoice(invoice_id):
    """Удалить наш неоплаченный счёт."""
    _call("crm.item.delete", {"entityTypeId": INVOICE_ETYPE, "id": invoice_id})


# ──────────────────────────── Сверка и применение ────────────────────────────

def _split_owned(deal_id):
    """Наши счета сделки → (live: {id: invoice}, paid_ids: set)."""
    owned = db.get_schedule_invoice_ids(deal_id)
    live = list_invoices_by_ids(owned)
    paid = {i for i, inv in live.items() if str(inv.get("stageId")) == STAGE_PAID}
    return owned, live, paid


def preview_schedule(deal_id, rows):
    """Что произойдёт при сохранении — без записи. Для окна подтверждения."""
    owned, live, paid = _split_owned(deal_id)
    referenced, to_create, to_update = set(), 0, 0
    for row in rows:
        inv = row.get("invoice_id")
        if inv and int(inv) in live:
            referenced.add(int(inv))
            if int(inv) not in paid:
                to_update += 1
        else:
            to_create += 1
    to_delete = [i for i in owned if i not in referenced and i not in paid]
    return {"create": to_create, "update": to_update,
            "delete": len(to_delete), "paid_kept": len(paid)}


def apply_schedule(deal_id, kind, total, rows, params=None):
    """Создать/обновить/удалить счета под новый график и сохранить снимок.

    rows: [{seq, due_date, amount, invoice_id?}] — полный график.
    Оплаченные счета неприкосновенны; лишние наши неоплаченные — удаляем.
    """
    ctx = deal_context(deal_id)
    owned, live, paid = _split_owned(deal_id)

    referenced = set()
    result_items, created, updated, deleted = [], [], [], []

    for row in rows:
        seq, amount, due = row["seq"], row["amount"], row["due_date"]
        inv = row.get("invoice_id")
        inv = int(inv) if inv else None
        if inv and inv in live:
            referenced.add(inv)
            if inv not in paid:
                update_invoice(inv, amount, due)
                updated.append(inv)
            result_items.append({"seq": seq, "due_date": due, "amount": amount, "invoice_id": inv})
        else:
            new_id = create_invoice(deal_id, seq, amount, due, ctx)
            created.append(new_id)
            result_items.append({"seq": seq, "due_date": due, "amount": amount, "invoice_id": new_id})

    # лишние наши счета (не в графике и не оплачены) — удаляем; оплаченные сохраняем в снимке
    for inv in owned:
        if inv in referenced or inv in paid:
            continue
        delete_invoice(inv)
        deleted.append(inv)

    db.save_payment_schedule(deal_id, kind, total, result_items, **(params or {}))
    return {"created": created, "updated": updated, "deleted": deleted, "items": result_items}


def get_state(deal_id):
    """Состояние для фронта: сохранённый график + сумма сделки + флаги оплаты по строкам."""
    ctx = deal_context(deal_id)
    saved = db.get_payment_schedule(deal_id)
    state = {
        "deal_id": deal_id,
        "deal_title": ctx["title"],
        "deal_sum": ctx["opportunity"],
        "currency": ctx["currency"],
        "schedule": None,
    }
    if not saved:
        return state

    live = list_invoices_by_ids([it["invoice_id"] for it in saved["items"] if it["invoice_id"]])
    items = []
    for it in saved["items"]:
        inv = live.get(it["invoice_id"]) if it["invoice_id"] else None
        items.append({
            "seq": it["seq"],
            "due_date": it["due_date"],
            "amount": it["amount"],
            "invoice_id": it["invoice_id"],
            "paid": bool(inv and str(inv.get("stageId")) == STAGE_PAID),
            "stage": inv.get("stageId") if inv else None,
        })
    state["schedule"] = {
        "kind": saved["kind"],
        "total": saved["total"],
        "advance": saved["advance"],
        "count": saved["count"],
        "freq_months": saved["freq_months"],
        "first_date": saved["first_date"],
        "items": items,
    }
    return state


# ──────────────────────────── Документ с таблицей платежей ────────────────────────────

def _fmt_money(v):
    """2165000.0 → '2 165 000,00' (пробелы разрядов, запятая-разделитель копеек)."""
    return f"{float(v or 0):,.2f}".replace(",", " ").replace(".", ",")


def _fmt_date_ru(iso):
    """'2026-08-01' → '01.08.2026'."""
    y, m, d = str(iso)[:10].split("-")
    return f"{d}.{m}.{y}"


def _inject_table(src_path, rows):
    """Вставляет таблицу платежей на место маркера [[PAYMENTS_TABLE]] в docx. Возвращает путь итога."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document(src_path)

    marker = None
    for p in doc.paragraphs:
        if PAYMENTS_MARKER in p.text:
            marker = p
            break

    tbl = doc.add_table(rows=1 + len(rows), cols=3)
    # рамки
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0"); el.set(qn("w:color"), "000000")
        borders.append(el)
    tbl._tbl.tblPr.append(borders)

    for c, txt in zip(tbl.rows[0].cells, ["№", "Срок оплаты", "Сумма, руб."]):
        c.paragraphs[0].add_run(txt).bold = True
    for i, row in enumerate(rows, 1):
        cs = tbl.rows[i].cells
        cs[0].paragraphs[0].add_run(str(row.get("seq", i)))
        cs[1].paragraphs[0].add_run(_fmt_date_ru(row["due_date"]))
        cs[2].paragraphs[0].add_run(_fmt_money(row["amount"]))

    if marker is not None:
        marker._p.addprevious(tbl._tbl)
        marker._p.getparent().remove(marker._p)

    out = src_path.replace(".docx", "_final.docx")
    doc.save(out)
    return out


def generate_document(deal_id, rows, total):
    """Формирует документ «График платежей» и прикрепляет его к сделке.

    Битрикс заполняет поля сделки/итог по шаблону и оставляет маркер [[PAYMENTS_TABLE]],
    затем мы сами вставляем таблицу платежей и кладём готовый файл в ленту сделки.
    """
    if not rows:
        raise RuntimeError("График пуст — нечего выводить в документ")

    # 1) документ Битрикса из шаблона (скаляры заполнит сам, таблицу оставит маркером)
    res = _call("crm.documentgenerator.document.add", {
        "templateId": PAYMENT_DOC_TEMPLATE_ID,
        "entityTypeId": 2,  # 2 = сделка
        "entityId": deal_id,
        "values": {"ScheduleTotal": _fmt_money(total)},
    }) or {}
    doc = res.get("document") or {}
    doc_id = doc.get("id")
    url = doc.get("downloadUrlMachine") or doc.get("downloadUrl")
    if not url:
        raise RuntimeError(f"Не удалось сгенерировать документ: {res}")

    # 2) скачать и вставить таблицу
    src = tempfile.mktemp(suffix=".docx")
    r = requests.get(url, timeout=30); r.raise_for_status()
    with open(src, "wb") as f:
        f.write(r.content)
    final = _inject_table(src, rows)

    # промежуточная запись генератора (с маркером вместо таблицы) больше не нужна
    if doc_id:
        try:
            _call("crm.documentgenerator.document.delete", {"id": doc_id})
        except Exception:
            pass

    # 3) прикрепить к сделке (лента)
    with open(final, "rb") as f:
        b64 = base64.b64encode(f.read()).decode()
    _call("crm.timeline.comment.add", {"fields": {
        "ENTITY_ID": deal_id, "ENTITY_TYPE": "deal",
        "COMMENT": "График платежей — сформированный документ.",
        "FILES": [["grafik_platezhey.docx", b64]],
    }})

    for pth in (src, final):
        try:
            os.remove(pth)
        except OSError:
            pass

    return {"document_id": doc_id}
